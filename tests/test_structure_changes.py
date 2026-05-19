from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from doccompare.comparison.structure_changes import detect_header_footer_changes


def _add_page_field(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_begin, instr, fld_sep, result, fld_end])


def _make_doc(path: Path, with_page_number: bool):
    doc = Document()
    doc.add_paragraph("Same body text.")
    if with_page_number:
        _add_page_field(doc.sections[0].footer.paragraphs[0])
    doc.save(path)


def test_detects_page_number_added_to_footer(tmp_path):
    original = tmp_path / "original.docx"
    modified = tmp_path / "modified.docx"
    _make_doc(original, with_page_number=False)
    _make_doc(modified, with_page_number=True)

    changes = detect_header_footer_changes(original, modified)

    assert len(changes) == 1
    assert changes[0]["kind"] == "footer"
    assert changes[0]["change_type"] == "added"
    assert "page number field added" in changes[0]["summary"]
