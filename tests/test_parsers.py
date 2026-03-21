import pytest
import docx
from docx.shared import Pt
from pathlib import Path
from doccompare.parsers.docx_parser import DocxParser
from doccompare.models import ElementType, TextFormatting


def test_docx_parser_extracts_headings(simple_original_docx):
    parser = DocxParser()
    doc = parser.parse(simple_original_docx)
    headings = [e for e in doc.elements if e.element_type == ElementType.HEADING]
    assert len(headings) >= 1
    assert headings[0].level == 1
    assert "Introduction" in headings[0].plain_text


def test_docx_parser_extracts_paragraphs(simple_original_docx):
    parser = DocxParser()
    doc = parser.parse(simple_original_docx)
    paras = [e for e in doc.elements if e.element_type == ElementType.PARAGRAPH]
    assert len(paras) >= 3


def test_docx_parser_extracts_formatting(fixtures_dir):
    path = fixtures_dir / "formatted.docx"
    d = docx.Document()
    p = d.add_paragraph()
    run = p.add_run("Bold text")
    run.bold = True
    run2 = p.add_run(" normal text")
    d.save(str(path))

    parser = DocxParser()
    doc = parser.parse(path)
    paras = [e for e in doc.elements if e.element_type == ElementType.PARAGRAPH]
    assert len(paras) >= 1
    bold_runs = [r for r in paras[0].runs if TextFormatting.BOLD in r.formatting]
    assert len(bold_runs) >= 1


def test_docx_parser_handles_tables(fixtures_dir):
    path = fixtures_dir / "with_table.docx"
    d = docx.Document()
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell A1"
    table.cell(0, 1).text = "Cell B1"
    table.cell(1, 0).text = "Cell A2"
    table.cell(1, 1).text = "Cell B2"
    d.save(str(path))

    parser = DocxParser()
    doc = parser.parse(path)
    rows = [e for e in doc.elements if e.element_type == ElementType.TABLE_ROW]
    assert len(rows) == 2
    assert len(rows[0].children) == 2


def test_docx_parser_supports_docx():
    parser = DocxParser()
    assert parser.supports(Path("test.docx"))
    assert not parser.supports(Path("test.pdf"))
