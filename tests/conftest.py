import pytest
from pathlib import Path
import docx
from docx.shared import Pt


@pytest.fixture(scope="session")
def fixtures_dir(tmp_path_factory):
    d = tmp_path_factory.mktemp("fixtures")
    return d


@pytest.fixture(scope="session")
def simple_original_docx(fixtures_dir):
    path = fixtures_dir / "original.docx"
    doc = docx.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the original document.")
    doc.add_paragraph("It contains several paragraphs.")
    doc.add_paragraph("This paragraph will be deleted.")
    doc.add_paragraph("Final paragraph.")
    doc.save(str(path))
    return path


@pytest.fixture(scope="session")
def simple_modified_docx(fixtures_dir):
    path = fixtures_dir / "modified.docx"
    doc = docx.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the modified document.")
    doc.add_paragraph("It contains several paragraphs.")
    doc.add_paragraph("This is a new paragraph added in the middle.")
    doc.add_paragraph("Final paragraph.")
    doc.save(str(path))
    return path


@pytest.fixture(scope="session")
def move_original_docx(fixtures_dir):
    path = fixtures_dir / "move_original.docx"
    doc = docx.Document()
    doc.add_paragraph("First paragraph stays here.")
    doc.add_paragraph("This block of text will be moved to a different location.")
    doc.add_paragraph("Third paragraph.")
    doc.save(str(path))
    return path


@pytest.fixture(scope="session")
def move_modified_docx(fixtures_dir):
    path = fixtures_dir / "move_modified.docx"
    doc = docx.Document()
    doc.add_paragraph("First paragraph stays here.")
    doc.add_paragraph("Third paragraph.")
    doc.add_paragraph("This block of text will be moved to a different location.")
    doc.save(str(path))
    return path
