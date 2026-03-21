import re
from pathlib import Path
from lxml import etree
import docx
import docx.text.paragraph
import docx.table
from docx.oxml.ns import qn
from doccompare.models import (
    ParsedDocument, DocumentElement, TextRun, ElementType, TextFormatting
)
from .base import DocumentParser


LIST_STYLES = {"List Bullet", "List Number", "List Paragraph", "List Continue"}

# Regex to detect heading styles like "Heading 1", "Heading 2 (Not in TOC)", etc.
_HEADING_RE = re.compile(r'^Heading\s+(\d+)', re.IGNORECASE)


def _to_pt(length):
    """Convert a docx Length to points, or None."""
    if length is None:
        return None
    try:
        return float(length.pt)
    except Exception:
        return None


def _effective_pf(para, attr):
    """Get effective paragraph format value by walking the style chain."""
    val = getattr(para.paragraph_format, attr, None)
    if val is not None:
        return val
    style = para.style
    while style:
        val = getattr(style.paragraph_format, attr, None)
        if val is not None:
            return val
        style = style.base_style
    return None


def _effective_font(run, para, attr):
    """Get effective font value by walking run style -> paragraph style chain."""
    val = getattr(run.font, attr, None)
    if val is not None:
        return val
    if run.style and run.style.font:
        val = getattr(run.style.font, attr, None)
        if val is not None:
            return val
    style = para.style
    while style:
        if style.font:
            val = getattr(style.font, attr, None)
            if val is not None:
                return val
        style = style.base_style
    return None


def _detect_heading_level_from_style_name(style_name: str) -> int:
    """Detect heading level from style name pattern like 'Heading 1 (Not in TOC)'."""
    m = _HEADING_RE.match(style_name)
    if m:
        return int(m.group(1))
    if style_name == "Section" or style_name == "Title":
        return 1
    return -1


class DocxParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def parse(self, file_path: Path) -> ParsedDocument:
        doc = docx.Document(str(file_path))
        elements = []
        para_idx = 0
        tbl_idx = 0

        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para = docx.text.paragraph.Paragraph(child, doc)
                elem = self._parse_paragraph(para, f"p_{para_idx}", doc)
                if elem is not None:
                    elements.append(elem)
                para_idx += 1
            elif tag == "tbl":
                tbl = docx.table.Table(child, doc)
                rows = self._parse_table(tbl, f"t_{tbl_idx}")
                elements.extend(rows)
                tbl_idx += 1

        metadata = {}
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author

        return ParsedDocument(elements=elements, metadata=metadata)

    def _get_list_info(self, doc, num_id: int, ilvl: int) -> dict:
        """Return num_fmt, lvl_text, and start_val for a numbering level.

        Follows numStyleLink -> styleLink chains to resolve indirect definitions.
        """
        result = {'num_fmt': 'bullet', 'lvl_text': '', 'start_val': 1}
        try:
            np = doc.part.numbering_part
            if np is None:
                return result
            nxml = np._element

            # Find the abstractNum for this numId
            abid = None
            for num in nxml.findall(qn('w:num')):
                if num.get(qn('w:numId')) == str(num_id):
                    abid_el = num.find(qn('w:abstractNumId'))
                    if abid_el is not None:
                        abid = abid_el.get(qn('w:val'))
                    break

            if abid is None:
                return result

            # Find the abstractNum element
            ab_elem = None
            for ab in nxml.findall(qn('w:abstractNum')):
                if ab.get(qn('w:abstractNumId')) == abid:
                    ab_elem = ab
                    break

            if ab_elem is None:
                return result

            # Follow numStyleLink if present (indirect definition)
            nsl = ab_elem.find(qn('w:numStyleLink'))
            if nsl is not None:
                link_name = nsl.get(qn('w:val'))
                # Find the abstractNum with matching styleLink
                for ab2 in nxml.findall(qn('w:abstractNum')):
                    sl = ab2.find(qn('w:styleLink'))
                    if sl is not None and sl.get(qn('w:val')) == link_name:
                        ab_elem = ab2
                        break

            # Extract level info
            for lvl in ab_elem.findall(qn('w:lvl')):
                if lvl.get(qn('w:ilvl')) == str(ilvl):
                    nf = lvl.find(qn('w:numFmt'))
                    lt = lvl.find(qn('w:lvlText'))
                    sv = lvl.find(qn('w:start'))
                    if nf is not None:
                        result['num_fmt'] = nf.get(qn('w:val'), 'bullet')
                    if lt is not None:
                        result['lvl_text'] = lt.get(qn('w:val'), '')
                    if sv is not None:
                        result['start_val'] = int(sv.get(qn('w:val'), '1'))
                    return result
        except Exception:
            pass
        return result

    def _parse_paragraph(self, para, element_id: str, doc=None) -> "DocumentElement | None":
        style_name = para.style.name if para.style else "Normal"

        # --- Numbering extraction ---
        # Walk paragraph XML then style chain to find numId and ilvl.
        # numId and ilvl may come from different levels in the style chain.
        list_style = ""
        list_numid = 0
        list_lvl_text = ""
        list_ilvl = 0
        has_numbering = False
        numbering_suppressed = False  # True if paragraph has explicit numId=0

        # 1) Check direct paragraph w:numPr
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId_el = numPr.find(qn('w:numId'))
                ilvl_el = numPr.find(qn('w:ilvl'))
                if numId_el is not None:
                    direct_numid = int(numId_el.get(qn('w:val'), 0))
                    if direct_numid == 0:
                        # Explicit numId=0 means "suppress numbering"
                        numbering_suppressed = True
                    else:
                        list_numid = direct_numid
                        has_numbering = True
                if ilvl_el is not None:
                    list_ilvl = int(ilvl_el.get(qn('w:val'), 0))

        # 2) Walk style chain for numId and ilvl (they may come from different levels)
        if not numbering_suppressed:
            found_numid = has_numbering
            found_ilvl = (pPr is not None and pPr.find(qn('w:numPr')) is not None
                          and pPr.find(qn('w:numPr')).find(qn('w:ilvl')) is not None)
            style = para.style
            while style and (not found_numid or not found_ilvl):
                try:
                    style_pPr = style.element.find(qn('w:pPr'))
                    if style_pPr is not None:
                        style_numPr = style_pPr.find(qn('w:numPr'))
                        if style_numPr is not None:
                            if not found_numid:
                                sni = style_numPr.find(qn('w:numId'))
                                if sni is not None:
                                    snumid = int(sni.get(qn('w:val'), 0))
                                    if snumid != 0:
                                        list_numid = snumid
                                        found_numid = True
                                        has_numbering = True
                            if not found_ilvl:
                                siv = style_numPr.find(qn('w:ilvl'))
                                if siv is not None:
                                    list_ilvl = int(siv.get(qn('w:val'), 0))
                                    found_ilvl = True
                except Exception:
                    pass
                style = style.base_style

        # Resolve numbering format
        if has_numbering and doc is not None:
            info = self._get_list_info(doc, list_numid, list_ilvl)
            list_style = info.get('num_fmt', 'bullet')
            list_lvl_text = info.get('lvl_text', '')

        # --- Element type detection ---
        heading_level = _detect_heading_level_from_style_name(style_name)
        if heading_level >= 1:
            elem_type = ElementType.HEADING
            level = heading_level
        elif has_numbering:
            elem_type = ElementType.LIST_ITEM
            level = list_ilvl
        elif any(s in style_name for s in LIST_STYLES) or "List" in style_name:
            elem_type = ElementType.LIST_ITEM
            level = 0
        else:
            elem_type = ElementType.PARAGRAPH
            level = 0

        # --- Run extraction with effective formatting ---
        runs = []
        for run in para.runs:
            if not run.text:
                continue
            fmt = set()
            if _effective_font(run, para, 'bold'):
                fmt.add(TextFormatting.BOLD)
            if _effective_font(run, para, 'italic'):
                fmt.add(TextFormatting.ITALIC)
            if run.underline:
                fmt.add(TextFormatting.UNDERLINE)
            if run.font.strike:
                fmt.add(TextFormatting.STRIKETHROUGH)

            font_name = _effective_font(run, para, 'name')
            font_size_obj = _effective_font(run, para, 'size')
            font_size = None
            if font_size_obj is not None:
                try:
                    font_size = float(font_size_obj.pt)
                except (AttributeError, TypeError):
                    try:
                        font_size = float(font_size_obj)
                    except (TypeError, ValueError):
                        pass

            runs.append(TextRun(
                text=run.text,
                formatting=fmt,
                font_name=font_name,
                font_size=font_size,
            ))

        # Skip completely empty paragraphs with no style significance
        if not runs and elem_type == ElementType.PARAGRAPH:
            return None

        # --- Paragraph-level formatting ---
        alignment = None
        try:
            align_val = _effective_pf(para, 'alignment')
            if align_val is not None:
                _align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
                alignment = _align_map.get(int(align_val), None)
        except (TypeError, ValueError):
            pass

        left_indent = _to_pt(_effective_pf(para, 'left_indent'))
        right_indent = _to_pt(_effective_pf(para, 'right_indent'))
        first_line_indent = _to_pt(_effective_pf(para, 'first_line_indent'))
        space_before = _to_pt(_effective_pf(para, 'space_before'))
        space_after = _to_pt(_effective_pf(para, 'space_after'))

        line_spacing = None
        try:
            ls = _effective_pf(para, 'line_spacing')
            if ls is not None:
                val = float(ls)
                if val > 10:
                    pt_val = _to_pt(ls)
                    line_spacing = pt_val / 12.0 if pt_val else None
                else:
                    line_spacing = val
        except (TypeError, ValueError):
            pass

        return DocumentElement(
            element_type=elem_type,
            runs=runs,
            level=level,
            element_id=element_id,
            list_style=list_style,
            list_numid=list_numid,
            list_lvl_text=list_lvl_text,
            list_ilvl=list_ilvl,
            alignment=alignment,
            left_indent_pt=left_indent,
            right_indent_pt=right_indent,
            first_line_indent_pt=first_line_indent,
            space_before_pt=space_before,
            space_after_pt=space_after,
            line_spacing=line_spacing,
        )

    def _parse_table(self, table, table_id: str) -> list:
        rows = []
        for r_idx, row in enumerate(table.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                cell_runs = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text:
                            continue
                        fmt = set()
                        if run.bold:
                            fmt.add(TextFormatting.BOLD)
                        if run.italic:
                            fmt.add(TextFormatting.ITALIC)
                        cell_runs.append(TextRun(text=run.text, formatting=fmt))
                cell_elem = DocumentElement(
                    element_type=ElementType.TABLE_CELL,
                    runs=cell_runs,
                    element_id=f"{table_id}_r{r_idx}_c{c_idx}",
                )
                cells.append(cell_elem)
            row_elem = DocumentElement(
                element_type=ElementType.TABLE_ROW,
                element_id=f"{table_id}_r{r_idx}",
                children=cells,
            )
            rows.append(row_elem)
        return rows
